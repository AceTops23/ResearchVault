# Standard library imports
import io
import os
import random
import re
import secrets
import PyPDF2
import nltk.data

# Third party imports
from bs4 import BeautifulSoup 
from collections import Counter
from docx import Document
from flask import Flask, jsonify, redirect, render_template, request, session, send_file, url_for, abort
from joblib import load
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from PyPDF2 import PdfFileReader, PdfFileWriter, PageObject
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from sklearn.preprocessing import MinMaxScaler
from sklearn.ensemble import RandomForestClassifier
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from transformers import BertTokenizer, BertForSequenceClassification
from werkzeug.utils import secure_filename

# Conditional imports (if any)
try:
    import openai
    import spacy
    import torch
    import torch.nn.functional as F
except ImportError:
    pass

# Local application imports
import db_connection
from db_connection import DBConnection

# Loading models and vectorizers
model = load('random_forest.joblib')
vectorizer = load('tfidf_vectorizer.joblib')

# Setting up BERT model and tokenizer
model_path = "Models\BERT Model" 
tokenizer = BertTokenizer.from_pretrained(model_path)
bert_model = BertForSequenceClassification.from_pretrained(model_path)

# Defining section names and order
section_names = {0: "Introduction", 1: "Method", 2: "Result", 3: "Discussion"}
section_order = ['Introduction', 'Method', 'Result', 'Discussion']

# Defining section keywords
section_keywords = {
    "Introduction": ["Introduction", "Background"],
    "Method": ["Methodology", "Software"],
    "Result": ["Result"],
    "Discussion": ["gathered", "researchers"]
}

# Setting up Flask app
app = Flask(__name__)
app.secret_key = secrets.token_hex(16)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx'}

# Loading Spacy model
nlp = spacy.load("en_core_web_sm")

# Setting up database connection
DB = 'database.db'
db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), DB)
db_connection = DBConnection(db_path)

# Setting OpenAI API key
openai.api_key = "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"


# Flask routes
@app.route('/')
def index():
    """Render the index page."""
    return render_template('index.html')


@app.route('/logout', methods=['POST'])
def logout():
    """Clear the session and return success status."""
    session.clear()
    return jsonify({'success': True})


@app.route('/forget-password')
def forget_password():
    """Render the forget password page."""
    return render_template('fp.html')


@app.route("/fromdocx")
def fromdocx():
    """Render the fromdocx page."""
    return render_template("fromdocx.html")


@app.route("/publish")
def publish():
    """Render the publish page."""
    return render_template("publish.html")


@app.route("/chatbot")
def chatbot():
    """Render the chatbot page."""
    return render_template("chatbot.html")


@app.route('/DV')
def DV():
    """Render the DV page."""
    return render_template('DV.html')

@app.route('/genimrad')
def genimrad():
    """Render the genimrad page."""
    return render_template('genimrad.html')

@app.route('/validate_login', methods=['POST'])
def validate_login():
    """Validate user login and set session state."""
    # Create a new database connection
    db = DBConnection(db_path)
    
    # Get the data from the request
    data = request.json
    email = data['email']
    password = data['password']
    
    # Validate the login credentials
    email_exists, password_match = db.validate_login(email, password)
    
    # Close the database connection
    db.close_connection()
    
    # Set the session state
    if email_exists and password_match:
        session['email'] = email
        isLoggedIn = True
    else:
        isLoggedIn = False
    
    # Prepare the response
    response = {
        'success': email_exists and password_match,
        'emailExists': email_exists,
        'passwordMatch': password_match,
        'isLoggedIn': isLoggedIn
    }
    
    # Return the response as JSON
    return jsonify(response)


@app.route('/session_state')
def session_state():
    """Check if user is logged in."""
    # Check if 'email' is in the session
    isLoggedIn = 'email' in session
    
    # Return the session state as JSON
    return jsonify({'isLoggedIn': isLoggedIn})


@app.route('/browse', methods=['GET'])
def browse():
    """Browse publications based on selected sort, field, year, and search query."""
    try:
        # Create a new database connection
        db_conn = DBConnection('database.db')
        
        # Get the parameters from the request
        selected_sort = request.args.get('sort', 'latest')
        selected_field = request.args.get('field', '')
        selected_year = request.args.get('year', '')
        search_query = request.args.get('search', '')
        
        # Fetch the publications from the database
        items, unique_subject_areas, unique_years = db_conn.fetch_publications(
            selected_sort, selected_field, selected_year, search_query)
        
        # Close the database connection
        db_conn.close_connection()
        
        # Render the browse page with the fetched publications and parameters
        return render_template('browse.html', items=items, subject_areas=unique_subject_areas,
                               unique_years=unique_years, selected_sort=selected_sort,
                               selected_field=selected_field, selected_year=selected_year,
                               search_query=search_query)
    
    except Exception as e:
        print("Error browsing publications:", e)
        
        # Render an error page if an exception occurs
        return render_template('404.html')


@app.route("/research")
def research():
    """Browse research publications based on search query."""
    try:
        # Create a new database connection
        db_conn = DBConnection('database.db')
        
        # Get the search query from the request
        search_query = request.args.get('search', '')
        
        # Fetch the research publications from the database
        items = db_conn.fetch_research_publications(search_query)
        
        # Close the database connection
        db_conn.close_connection()
        
        # Render the research page with the fetched publications and search query
        return render_template('research.html', items=items, search_query=search_query)
    
    except Exception as e:
        print("Error browsing publications:", e)
        
        # Render an error page if an exception occurs
        return render_template('404.html')


@app.route("/api", methods=["POST"])
def api():
    """Handle POST requests to the /api route."""
    # Get the message from the request
    message = request.json.get("message")
    
    # Use OpenAI's GPT-3 model to generate a response
    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": message}
        ]
    )
    
    # If a message was generated, return it; otherwise, return an error message
    if completion.choices[0].message is not None:
        return completion.choices[0].message
    else:
        return 'Failed to Generate response!'


def save_file(file):
    """Save a file and return its filename."""
    if file:
        filename = file.filename
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        return filename
    return None


@app.route('/submit_data', methods=['POST'])
def submit_data():
    """Handle POST requests to the /submit_data route."""
    # Get the form data from the request
    title = request.form['title']
    authors = request.form['authors']
    publicationDate = request.form['publicationDate']
    thesisAdvisor = request.form['thesisAdvisor']
    department = request.form['department']
    degree = request.form['degree']
    subjectArea = request.form['subjectArea']
    abstract = request.form['abstract']
    
    # Save the uploaded file and get its filename
    uploaded_file = request.files['file']
    filename = save_file(uploaded_file)
    
    # If a file was uploaded, insert a new record into the database
    if filename:
        if db_connection.insert_upload(title, authors, publicationDate, thesisAdvisor, department, degree, subjectArea, abstract, filename):
            print("Upload record inserted successfully!")
            return "Upload successful!"
        else:
            print("Failed to insert upload record.")
    
    return "Data submitted successfully!"


@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle POST requests to the /upload route."""
    try:
        # Get the form data from the request
        title = request.form.get('title')
        file = request.files.get('file')
        
        # Validate the form data
        if not title or not file:
            return 'Please fill in all fields.', 400
        if not file.filename.endswith('.docx'):
            return 'Invalid file type. Please upload a .docx file.', 400
        
        # Save the uploaded file and get its filename
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Create a new database connection and insert a new record into the database
        db_conn = DBConnection('database.db')
        
        if filename:
            if db_conn.insert_into_working(title, file_path):
                print("Upload record inserted successfully!")
                return "Upload successful!"
            else:
                print("Failed to insert upload record.")
                return "Upload failed.", 500
        
        db_conn.close_connection()
    
    except Exception as e:
        print(f"An error occurred: {e}")
        
        # Return an error message if an exception occurs
        return "An error occurred while processing your request.", 500
    
    return "Data submitted successfully!"


@app.route('/abstract')
def abstract():
    """Render the abstract page for the last unapproved record."""
    # Get the last unapproved record from the database
    record = db_connection.get_last_unapproved()
    
    # If a record was found, render the abstract page with the record's title
    if record is not None:
        title = record['title']
        return render_template('abstract.html', title=title)
    else:
        return "No unapproved records found."
    
    
@app.route('/get_last_unapproved')
def get_last_unapproved_route():
    """Return the last unapproved record as JSON."""
    # Get the last unapproved record from the database
    record = db_connection.get_last_unapproved()
    
    # If there's no record, return an error
    if record is None:
        abort(500, 'No records in database')
    
    # Return the record as JSON
    return jsonify(record)


@app.route('/upload_abstract', methods=['POST'])
def upload_abstract_route():
    """Update the abstract of the last unapproved record."""
    # Get the abstract from the request data
    data = request.get_json()
    abstract = data['abstract']
    
    # Get the last unapproved record from the database
    record = db_connection.get_last_unapproved()
    
    # If a record was found, update its abstract in the database
    if record:
        success = db_connection.update_abstract(record['id'], abstract)
        
        # Return a success status if the update was successful
        if success:
            return jsonify({'status': 'success'})
    
    # Return a failure status if no record was found or the update failed
    return jsonify({'status': 'failure'})


@app.route('/generate_abstract')
def generate_abstract():
    """Generate an abstract for the last unapproved record."""
    print("Starting to generate abstract...")
    
    # Get the last unapproved record from the database
    record = db_connection.get_last_unapproved()
    
    if record is not None:
        print("Record found, processing...")
        
        # Get the path of the document file
        file_path = record['IMRAD'] 
        
        # Read the document and extract its text
        doc = Document(file_path)
        doc_text = " ".join([p.text for p in doc.paragraphs])
        
        # Clean up the text
        soup = BeautifulSoup(doc_text, 'html.parser')
        cleaned_text = re.sub(r'\s+', ' ', soup.get_text(separator=' '))
        
        # Split the text into sentences
        sentences = sent_tokenize(cleaned_text)
        
        # Split the sentences into chunks of a certain size
        chunk_size = 512
        text_chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += " " + sentence
            else:
                text_chunks.append(current_chunk)
                current_chunk = sentence
        
        if current_chunk:
            text_chunks.append(current_chunk)
        
        # Classify each chunk into a section using a BERT model
        section_texts = {section: [] for section in section_names.values()}
        section_texts['Other'] = []
        
        batch_size = 20
        
        for i in range(0, len(text_chunks), batch_size):
            print(f"Classifying batch {i//batch_size + 1}...")
            
            batch = text_chunks[i:i+batch_size]
            inputs = tokenizer(batch, return_tensors="pt", padding=True, truncation=True)
            
            with torch.no_grad():
                outputs = bert_model(**inputs)
                predicted_sections = torch.argmax(outputs.logits, dim=1)
            
            for j, input_ids in enumerate(inputs["input_ids"]):
                token_text = tokenizer.decode(input_ids, skip_special_tokens=True, clean_up_tokenization_spaces=True)
                current_section = section_names[predicted_sections[j].item()]
                
                for section, keywords in section_keywords.items():
                    if any(keyword in token_text for keyword in keywords):
                        current_section = section
                
                section_texts[current_section].append((token_text, outputs.logits[j]))
        
        # Select the best chunks from each section and calculate their average probability
        total_probability = 0
        total_chunks = 0
        sorted_section_texts = {}
        
        selected_chunks = []
        
        for section in section_order:
            if section in section_texts:
                print(f"Processing section {section}...")
                
                top_chunks = sorted(section_texts[section], key=lambda x: F.softmax(x[1], dim=0).max().item(), reverse=True)[:10]
                selected_chunk, logits = random.choice([chunk for chunk in top_chunks if chunk not in selected_chunks])
                
                selected_chunks.append(selected_chunk)
                
                sorted_section_texts[section] = f"\n\n{selected_chunk}"
                
                total_probability += F.softmax(logits, dim=0).max().item()
                total_chunks += 1
        
        average_probability = total_probability / total_chunks if total_chunks > 0 else 0
        
        print(f"Average probability of chosen chunks: {average_probability}")
        
        print("Finished processing. Sending response...")
    
    return jsonify({'section_texts': sorted_section_texts, 'average_probability': average_probability})


@app.route('/publication_detail/<int:item_id>')
def publication_detail(item_id):
    """Render the publication detail page for a specific record."""
    # Create a new database connection and get the publication by its ID
    db_conn = DBConnection(DB)
    item = db_conn.get_publication_by_id(item_id)
    db_conn.close_connection()
    
    # If a record was found, read its PDF file and render the publication detail page
    if item is not None:
        pdf_path = os.path.join('uploads', item['file_path'])
        text_content = read_pdf_text(pdf_path)
        return render_template('publication_detail.html', item=item, text_content=text_content)
    else:
        return render_template('404.html'), 404


def read_pdf_text(pdf_path):
    """Read the text from a PDF file."""
    text_content = ""
    
    # Open the PDF file and create a reader object
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        
        # Read the text from each page of the PDF
        num_pages = len(pdf_reader.pages)
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text_content += page.extract_text()
    
    return text_content


@app.route('/uploads/<path:filename>')
def serve_pdf(filename):
    """Serve a PDF file."""
    pdf_path = os.path.join('uploads', filename)
    return send_file(pdf_path, mimetype='application/pdf')


def generate_apa_citation_from_data(publication):
    """Generate an APA citation for a publication."""
    authors = publication['authors'].split('; ')
    
    # Format the authors depending on how many there are
    num_authors = len(authors)
    if num_authors == 1:
        formatted_authors = authors[0].split()[-1]
    elif num_authors == 2:
        formatted_authors = f"{authors[0].split()[-1]} & {authors[1].split()[-1]}"
    else:
        formatted_authors = ", ".join(author.split()[-1] for author in authors[:-1])
        formatted_authors += f", & {authors[-1].split()[-1]}"
    
    # Generate the APA citation
    apa_citation = f"{formatted_authors}. ({publication['year']}). {publication['title']}. {publication['thesisAdvisor']}. {publication['department']}. {publication['degree']}."
    
    return apa_citation


@app.route('/generate_apa_citation/<int:item_id>')
def generate_apa_citation(item_id):
    """Generate an APA citation for a specific record."""
    # Create a new database connection and get the publication by its ID
    db_conn = DBConnection(DB)
    publication = db_conn.get_publication_by_id(item_id)
    
    # If a record was found, generate an APA citation for it
    if publication:
        apa_citation = generate_apa_citation_from_data(publication)
        return jsonify({"apa_citation": apa_citation})
    
    return jsonify({"error": "Publication not found"})


def simpleSplit(text, fontName, fontSize, maxWidth):
    """Split a text into lines based on a maximum width."""
    words = text.split(' ')
    
    lines = []
    currentLine = ''
    
    # Add words to the current line until it exceeds the maximum width
    for word in words:
        if stringWidth(currentLine + ' ' + word, fontName, fontSize) <= maxWidth:
            currentLine += ' ' + word
        else:
            lines.append(currentLine)
            currentLine = word
    
    if currentLine:
        lines.append(currentLine)
    
    return lines


def clean_text(text):
    """Clean up a text by removing non-word characters and converting to lowercase."""
    # Remove non-word characters
    text = re.sub(r'\W+', ' ', text)
    
    # Convert to lowercase
    text = text.lower()
    
    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    word_tokens = word_tokenize(text)
    text = [word for word in word_tokens if not word in stop_words]
    
    return ' '.join(text)


def convert_to_imrad(file_path):
    """Convert a PDF file to IMRaD format."""
    try:
        # Open the PDF file and extract its text
        with open(file_path, "rb") as f:
            pdf_reader = PdfFileReader(f)
            pdf_text = ""
            for page in pdf_reader.pages:
                pdf_text += page.extract_text()
        
        # Clean up the text
        pdf_text = clean_text(pdf_text)
        
        # Split the text into chunks of a certain size
        chunk_size = 512
        text_chunks = [pdf_text[i:i+chunk_size] for i in range(0, len(pdf_text), chunk_size)]
        
        # Classify each chunk into a section using a BERT model
        section_texts = {section: "" for section in section_names.values()}
        
        batch_size = 10
        
        for i in range(0, len(text_chunks), batch_size):
            batch = text_chunks[i:i+batch_size]
            inputs = tokenizer(batch, return_tensors="pt", padding=True, truncation=True)
            
            with torch.no_grad():
                outputs = bert_model(**inputs)
                predicted_sections = torch.argmax(outputs.logits, dim=1)
            
            for j, input_ids in enumerate(inputs["input_ids"]):
                current_section = section_names[predicted_sections[j].item()]
                
                for token in input_ids:
                    token_text = tokenizer.decode([token.item()], skip_special_tokens=True, clean_up_tokenization_spaces=True)
                    
                    if current_section is not None:
                        if section_texts[current_section] and not section_texts[current_section].endswith(' '):
                            section_texts[current_section] += ' '
                        section_texts[current_section] += token_text
                
                print(f"Section: {current_section}")
                print(f"Text: {section_texts[current_section]}")
                print(f"Token: {token_text}")
        
        # Clean up the section texts and remove unnecessary spaces
        section_texts[current_section] = clean_text(section_texts[current_section])            
        section_texts[current_section] = section_texts[current_section].replace(' ##', '')                               
        
        # Convert the section texts into a new PDF file
        converted_file_path = file_path.replace(os.path.splitext(file_path)[1], '_imrad.pdf')
        
        pdf_writer = PdfFileWriter()
        
        for section_name, section_text in section_texts.items():
            page = PageObject.createBlankPage(None, 595.44, 841.68)
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=A4)
            
            textobject = can.beginText()
            textobject.setTextOrigin(72, 841.68 - 72)
            textobject.setFont("Helvetica", 10)
            textobject.textLine(section_name)
            
            for line in section_text.split('\n'):
                lines = simpleSplit(line, "Helvetica", 10, 595.44 - 144)
                
                for l in lines:
                    textobject.textLine(l)
            
            can.drawText(textobject)
            can.save()
            
            packet.seek(0)
            new_pdf = PdfFileReader(packet)
            
            page.mergePage(new_pdf.getPage(0))
            
            pdf_writer.addPage(page)
        
        with open(converted_file_path, "wb") as out_f:
            pdf_writer.write(out_f)
        
        print(f"Converted file saved at {converted_file_path}")
        
        return converted_file_path
    
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        
        
@app.route('/convert_to_imrad/<int:item_id>', methods=['GET'])
def convert_to_imrad_route(item_id):
    """Convert a specific publication to IMRaD format."""
    # Get the publication by its ID
    publication = db_connection.get_publication_by_id(item_id)
    
    # If no publication was found, return an error message
    if publication is None:
        return jsonify({"message": "Publication not found."}), 404
    
    # Get the path of the publication file
    file_name = os.path.basename(publication['file_path'])
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_name)
    
    try:
        # Convert the publication to IMRaD format and update its path in the database
        converted_file_path = convert_to_imrad(file_path)
        db_connection.update_converted_file_path(item_id, converted_file_path)
        
        # Return the converted file
        return send_file(converted_file_path, as_attachment=False)
    
    except Exception as e:
        # Return an error message if an exception occurs
        return jsonify({"message": "Error converting to IMRAD.", "error": str(e)}), 500


@app.route('/convert_docx_to_imrad', methods=['POST'])
def convert_docx_to_imrad_route():
    """
    This route handles POST requests to convert DOCX files to IMRaD format."""
    # Get the JSON data from the request
    data = request.get_json()
    
    # Extract the file path from the data
    file_path = data['file_path']
    
    # Call the function to convert the DOCX file to IMRaD format
    converted_file_path = convert_docx_to_imrad(file_path)
    
    # Update the database with the path of the converted file
    db_connection.update_imrad_path(file_path, converted_file_path)
    
    # Return the path of the converted file in a JSON object
    return jsonify({'converted_file_path': converted_file_path})


def convert_docx_to_imrad(file_path):
    """Convert a DOCX file to IMRaD format."""
    try:
        # Open the DOCX file and extract its text
        doc = Document(file_path)
        doc_text = "\n".join([para.text for para in doc.paragraphs])
        
        # Clean up the text
        doc_text = clean_text(doc_text)
        
        # Split the text into chunks of a certain size with overlap
        chunk_size = 512
        overlap = 50
        text_chunks = [doc_text[i:i+chunk_size] for i in range(0, len(doc_text), chunk_size-overlap)]
        
        # Convert the chunks into a bag-of-words representation
        vectorizer = CountVectorizer()
        X = vectorizer.fit_transform(text_chunks)

        # Normalize the data to a range between 0 and 1
        scaler = MinMaxScaler()
        X_normalized = scaler.fit_transform(X.toarray())

        # Classify each chunk into a section using a BERT model
        section_texts = {section: "" for section in section_names.values()}
        
        batch_size = 10
        threshold = 0.5  # Set the threshold
        
        for i in range(0, len(text_chunks), batch_size):
            batch = text_chunks[i:i+batch_size]
            inputs = tokenizer(batch, return_tensors="pt", padding=True, truncation=True)
            
            with torch.no_grad():
                outputs = bert_model(**inputs)
                predicted_sections = torch.argmax(outputs.logits, dim=1)
            
            for j, input_ids in enumerate(inputs["input_ids"]):
                current_section = section_names.get(predicted_sections[j].item(), 'Other')
                
                for k, token in enumerate(input_ids):
                    token_text = tokenizer.decode([token.item()], skip_special_tokens=True, clean_up_tokenization_spaces=True)
                    max_prob = F.softmax(outputs.logits[j], dim=0).max().item()
                    
                    if k == 0 and token_text.startswith('##'):
                        # If the first token is an incomplete word (starts with '##'), look back at the previous chunk to find the complete word
                        prev_chunk_last_word = tokenizer.decode([inputs["input_ids"][j-1][-1].item()], skip_special_tokens=True, clean_up_tokenization_spaces=True)
                        token_text = prev_chunk_last_word + token_text[2:]
                    
                    if current_section is not None and max_prob >= threshold:
                        if section_texts[current_section] and not section_texts[current_section].endswith(' '):
                            section_texts[current_section] += ' '
                        section_texts[current_section] += token_text
                
                print(f"Section: {current_section}")
                print(f"Text: {section_texts[current_section]}")
                print(f"Token: {token_text}")
        
        # Clean up the section texts and remove unnecessary spaces
        for current_section in section_texts:
            section_texts[current_section] = clean_text(section_texts[current_section])                              
        
        # Convert the section texts into a new DOCX file
        converted_file_path = file_path.replace(os.path.splitext(file_path)[1], '_imrad.docx')
        
        converted_doc = Document()
        
        for section_name, section_text in section_texts.items():
            converted_doc.add_heading(section_name, level=1)
            converted_doc.add_paragraph(section_text)
        
        converted_doc.save(converted_file_path)
        
        print(f"Converted file saved at {converted_file_path}")
        
        return converted_file_path
    
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        

@app.route('/convert_docx_to_text', methods=['POST'])
def convert_docx_to_text():
    data = request.get_json()
    file_path = data['file_path']

    # Open the DOCX file and extract its text
    doc = Document(file_path)
    doc_text = "\n".join([para.text for para in doc.paragraphs])

    return jsonify({'text_content': doc_text})
        

if __name__ == '__main__':
    """Run the Flask app."""
    app.run()

